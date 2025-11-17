"""
Utility functions for document loading, cleaning, and processing.
"""

import re
from pathlib import Path
from typing import List, Tuple
import fitz  # PyMuPDF for faster PDF extraction
import docx


class DocumentLoader:
    """Load and extract text from various document formats."""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> Tuple[str, str]:
        """
        Extract text from PDF file using PyMuPDF (faster than PyPDF2).
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Tuple of (text, metadata)
        """
        text = ""
        metadata = ""
        try:
            doc = fitz.open(file_path)
            metadata = f"Pages: {len(doc)}"
            
            for page_num, page in enumerate(doc, 1):
                text += f"--- Page {page_num} ---\n"
                text += page.get_text("text") + "\n"
            
            doc.close()
        except Exception as e:
            raise ValueError(f"Error reading PDF: {str(e)}")
        
        return text, metadata
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> Tuple[str, str]:
        """Extract text from TXT file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            metadata = f"Characters: {len(text)}"
            return text, metadata
        except Exception as e:
            raise ValueError(f"Error reading TXT: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> Tuple[str, str]:
        """Extract text from DOCX file."""
        text = ""
        try:
            doc = docx.Document(file_path)
            paragraphs = len(doc.paragraphs)
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            metadata = f"Paragraphs: {paragraphs}"
            return text, metadata
        except Exception as e:
            raise ValueError(f"Error reading DOCX: {str(e)}")
    
    @staticmethod
    def extract_text(file_path: str) -> Tuple[str, str]:
        """
        Extract text from any supported file format.
        
        Args:
            file_path: Path to file
            
        Returns:
            Tuple of (text, metadata)
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == ".pdf":
            return DocumentLoader.extract_text_from_pdf(file_path)
        elif file_ext == ".txt":
            return DocumentLoader.extract_text_from_txt(file_path)
        elif file_ext == ".docx":
            return DocumentLoader.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")


class TextCleaner:
    """Clean and normalize text."""
    
    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean text by removing extra whitespace and special characters.
        
        Args:
            text: Raw text
            
        Returns:
            Cleaned text
        """
        # Remove multiple spaces/newlines
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special Unicode characters but keep basic punctuation
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)\[\]\"\']', '', text)
        
        # Trim
        text = text.strip()
        
        return text
    
    @staticmethod
    def extract_sentences(text: str) -> List[str]:
        """
        Extract sentences from text.
        
        Args:
            text: Raw text
            
        Returns:
            List of sentences
        """
        # Split by period, exclamation, question mark
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]
    
    @staticmethod
    def extract_key_terms(text: str, top_n: int = 10) -> List[str]:
        """
        Extract key terms from text (simple frequency-based).
        
        Args:
            text: Raw text
            top_n: Number of top terms to extract
            
        Returns:
            List of key terms
        """
        # Extract words
        words = re.findall(r'\b[a-z]+\b', text.lower())
        
        # Filter stop words
        stop_words = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
            'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are', 'be', 'been',
            'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would',
            'could', 'should', 'may', 'might', 'must', 'can', 'this', 'that',
            'these', 'those', 'i', 'you', 'he', 'she', 'it', 'we', 'they'
        }
        
        words = [w for w in words if w not in stop_words and len(w) > 2]
        
        # Count frequency
        from collections import Counter
        freq = Counter(words)
        
        # Return top N
        return [word for word, _ in freq.most_common(top_n)]


class TextChunker:
    """Split text into chunks with overlap."""
    
    @staticmethod
    def chunk_text(
        text: str,
        chunk_size: int = 1000,
        overlap: int = 200
    ) -> List[Tuple[str, int]]:
        """
        Split text into overlapping chunks.
        
        Args:
            text: Text to chunk
            chunk_size: Size of each chunk in characters (default 1000)
            overlap: Overlap between chunks (default 200)
            
        Returns:
            List of (chunk_text, chunk_start_position) tuples
        """
        # Clean text first
        text = TextCleaner.clean_text(text)
        
        if len(text) <= chunk_size:
            return [(text, 0)]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            chunks.append((chunk, start))
            start += (chunk_size - overlap)
        
        return chunks
    
    @staticmethod
    def chunk_by_sentences(
        text: str,
        sentences_per_chunk: int = 5,
        overlap_sentences: int = 1
    ) -> List[str]:
        """
        Chunk text by sentences instead of fixed size.
        
        Args:
            text: Text to chunk
            sentences_per_chunk: Number of sentences per chunk
            overlap_sentences: Number of overlapping sentences
            
        Returns:
            List of chunked text
        """
        sentences = TextCleaner.extract_sentences(text)
        
        if len(sentences) <= sentences_per_chunk:
            return [' '.join(sentences)]
        
        chunks = []
        start = 0
        
        while start < len(sentences):
            end = start + sentences_per_chunk
            chunk = ' '.join(sentences[start:end])
            chunks.append(chunk)
            start += (sentences_per_chunk - overlap_sentences)
        
        return chunks
